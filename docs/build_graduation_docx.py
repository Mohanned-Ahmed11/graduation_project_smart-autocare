#!/usr/bin/env python3
"""
Build Smart_Auto_Graduation_Documentation.docx with logo + UML PNGs + explanations.
Requires: pip install python-docx
Run from repo root: python docs/build_graduation_docx.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
PLANTUML = DOCS / "plantuml"
LOGO = ROOT / "assets" / "logo.png"
OUT = DOCS / "Smart_Auto_Graduation_Documentation.docx"


def add_paragraphs(doc: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        doc.add_paragraph(text)


# (image path relative to PLANTUML or ROOT, title, explanation paragraphs)
SECTIONS: list[tuple[Path, str, list[str]]] = [
    (
        PLANTUML / "SmartAutoCar_ERD.png",
        "1. Entity–Relationship Diagram (ERD)",
        [
            "This diagram shows the main database tables for Smart Auto and how they connect.",
            "Each person who signs up gets a row in profiles (linked to Supabase auth.users). "
            "issues stores AI diagnosis runs (text, photo, voice links, and the AI JSON answer). "
            "requests are “help me” posts. request_responses links a driver to a request (pending or accepted). "
            "chats has one chat per request, and messages stores each chat line with sender and optional read time.",
            "The lines between boxes are foreign keys: data stays consistent when rows are deleted (cascade rules in SQL).",
        ],
    ),
    (
        PLANTUML / "SmartAutoCar_UseCase.png",
        "2. Use Case Diagram",
        [
            "This diagram shows who uses the app and what they can do in simple blocks.",
            "Guests can only reach sign-in. Logged-in users can manage their profile, create AI issues, create and browse help requests, respond to requests, chat, and receive realtime updates.",
            "Supabase appears as security (RLS) and file storage because those features sit behind the scenes for almost every action.",
        ],
    ),
    (
        PLANTUML / "SmartAutoCar_Activity.png",
        "3. Activity Diagram (help request flow)",
        [
            "This is the step-by-step flow for the “community help” feature.",
            "The requester registers, gets a profile automatically, and creates an open request. A driver lists open requests and sends a response. "
            "The requester accepts one driver; then the system can open a chat tied to that request. Both sides can send messages.",
            "Notes on the diagram mention database rules (RLS) and realtime messages.",
        ],
    ),
    (
        PLANTUML / "SmartAutoCar_Block.png",
        "4. Block Diagram (system overview)",
        [
            "This diagram is a high-level map of the whole system—not code detail, just the big boxes.",
            "The Flutter app talks to Supabase for login, database (PostgreSQL), file storage, and realtime channels. "
            "It also talks to Maps for location and to an AI service (e.g. Gemini) for diagnosis.",
            "Dotted lines show logical links: auth syncs with profile data, the database stores file URLs, and realtime pushes live updates.",
        ],
    ),
    (
        PLANTUML / "sequence_diagram_auth_profile.png",
        "5. Sequence Diagram — Sign up and profile",
        [
            "This diagram shows the order of messages when a new user registers.",
            "The app calls Supabase Auth. When auth.users gets a new row, a database trigger runs and inserts a matching profiles row (function handle_new_user).",
            "Later, the user can update their profile; the database only allows changes to their own row (RLS).",
        ],
    ),
    (
        PLANTUML / "sequence_diagram_help_request_chat.png",
        "6. Sequence Diagram — Request, accept, and chat",
        [
            "This diagram follows time from top to bottom for the help-request and chat scenario.",
            "The requester creates a request; the driver reads open requests and inserts a response. The requester accepts; both update request and response status.",
            "A chat row is tied to the request. When either side inserts a message, Supabase Realtime can notify the other user so the UI updates quickly.",
        ],
    ),
    (
        PLANTUML / "sequence_diagram_ai_issue_storage.png",
        "7. Sequence Diagram — AI issue and storage",
        [
            "This diagram shows how an AI diagnosis issue is stored.",
            "The user picks a photo or records voice. The app uploads files to the correct storage bucket (paths scoped by user id in the project rules).",
            "The app inserts a row in issues with the file URLs, calls the AI API, then saves the structured answer in the ai_response JSON field.",
        ],
    ),
]


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Smart Auto", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "Graduation project — full description, related apps, and UML diagrams"
    )
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(1.6))
        doc.add_paragraph()
    else:
        doc.add_paragraph(f"[Logo not found at {LOGO}]")

    # ----- Introduction -----
    doc.add_heading("Introduction", level=1)
    add_paragraphs(
        doc,
        [
            "Smart Auto is a cross-platform mobile application built as a graduation project. It targets everyday drivers "
            "who want quick guidance when something seems wrong with their car, who need to find nearby automotive "
            "services, and who may want to ask other people for practical help or a fair price for a small job.",
            "The project combines a modern client (Flutter) with a cloud backend (Supabase) so that user accounts, "
            "stored diagnoses, help requests, and chat messages stay secure, consistent, and available from different "
            "devices. Optional integrations such as maps and generative AI extend the app beyond a simple static form.",
            "This document gives an introduction, a clear problem statement, a list of features, a comparison with related "
            "apps, a detailed explanation of how Smart Auto works, and UML diagrams that match the database design used in "
            "the Supabase SQL migrations.",
        ],
    )

    # ----- Problem definition -----
    doc.add_heading("Problem definition", level=1)
    add_paragraphs(
        doc,
        [
            "Many drivers are not experts in vehicle mechanics. When a warning light appears, a noise starts, or "
            "performance drops, they often feel uncertain about how serious the problem is and what to do next. "
            "Searching the web can return confusing or contradictory advice, and visiting a workshop immediately may "
            "be costly or inconvenient if the issue is minor.",
            "Finding a trustworthy garage or service nearby is another pain point. People may not know which businesses "
            "are close, open, or relevant to their problem. Paper maps and word of mouth are less common today; users "
            "expect map-based discovery on the phone.",
            "Finally, peer help is underused: a driver nearby might be willing to advise, lend a tool, or agree on a "
            "small paid task, but there is usually no safe, structured channel to post a request, receive responses, "
            "choose someone, and chat with clear boundaries.",
            "Smart Auto is proposed as a single place that addresses these gaps: understandable AI-assisted diagnosis "
            "(with the limitation that it is not a substitute for a professional inspection), map-oriented discovery of "
            "services, and a moderated workflow for community help requests with chat tied to an accepted helper.",
        ],
    )

    # ----- Features -----
    doc.add_heading("Main features of the application", level=1)
    features = [
        (
            "Onboarding and branding",
            "Guided first screens introduce the product; the app identity uses the Smart Auto logo and a consistent theme.",
        ),
        (
            "Authentication and profile",
            "Users sign up and sign in through Supabase Auth. A profile row is created automatically and can store name, "
            "phone, car model, profile image, and last known location for map-related features.",
        ),
        (
            "AI-assisted diagnosis (issues)",
            "Users can describe a problem using text, photo, and/or voice. Files are stored in protected buckets; the app "
            "calls an AI model (e.g. Gemini) and saves structured results in the database for later review.",
        ),
        (
            "Home and navigation",
            "A central home area gives quick access to diagnosis, requests, maps, and settings (including light/dark theme).",
        ),
        (
            "Help requests and responses",
            "A user can publish a help request with title, description, optional image, suggested price, and location. "
            "Other authenticated users can browse open requests and submit a response. The owner can accept one helper; "
            "status fields follow the schema (open, accepted, completed, etc.).",
        ),
        (
            "Realtime chat",
            "Each accepted help flow can use a dedicated chat and message history. Supabase Realtime keeps new messages "
            "visible without constant manual refresh.",
        ),
        (
            "Maps and nearby services",
            "Integration with maps and location helps users see where they are and find nearby car-related points of interest.",
        ),
        (
            "Security at the data layer",
            "Row Level Security (RLS) policies ensure users only access profiles, issues, requests, and messages they are "
            "allowed to see, according to the rules defined in the SQL script.",
        ),
    ]
    for title, desc in features:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(title + ". ").bold = True
        p.add_run(desc)

    # ----- Related applications -----
    doc.add_heading("Related and similar applications", level=1)
    add_paragraphs(
        doc,
        [
            "Several categories of existing products overlap partly with Smart Auto. None of them alone matches the full "
            "combination of AI diagnosis, open help requests with acceptance, realtime chat, and Supabase-backed rules "
            "as implemented in this graduation project.",
        ],
    )
    related = [
        (
            "Google Maps / Apple Maps",
            "Strong for navigation and finding businesses, but not designed for AI car diagnosis or structured peer help "
            "requests with per-request chat.",
        ),
        (
            "Waze",
            "Community reports and navigation; social layer is traffic-focused, not mechanical diagnosis or paid micro-help.",
        ),
        (
            "RepairPal, YourMechanic, or similar",
            "Often oriented toward booking certified mechanics or price estimates. Smart Auto is lighter and student-built, "
            "with emphasis on AI explanation plus optional community responses rather than a full marketplace.",
        ),
        (
            "OBD / scanner apps (e.g. Torque-style tools)",
            "Read live vehicle data from hardware. Smart Auto does not require an OBD dongle in the MVP; it focuses on "
            "symptom description and media instead.",
        ),
        (
            "Generic messaging (WhatsApp, Telegram)",
            "Good for free-form chat but lack request lifecycle, database-enforced participation, and integrated maps/diagnosis.",
        ),
    ]
    for name, text in related:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(name + ". ").bold = True
        p.add_run(text)

    # ----- Detailed app explanation -----
    doc.add_heading("Detailed description of the Smart Auto application", level=1)

    doc.add_heading("User journey (high level)", level=2)
    add_paragraphs(
        doc,
        [
            "A new user opens the app, sees onboarding, then registers or logs in. After authentication, the home screen "
            "becomes the hub. From there the user may run an AI diagnosis flow (create an issue record, upload evidence, "
            "read the AI output), browse or create help requests, open the map for nearby services, or adjust profile "
            "and appearance settings.",
            "When two users connect through a help request, the system moves from a public or semi-public request list to a "
            "private chat channel bound to that request, so conversation context stays clear.",
        ],
    )

    doc.add_heading("Core modules in the codebase", level=2)
    add_paragraphs(
        doc,
        [
            "The Flutter project is organised by feature areas (for example onboarding, authentication, home, AI detection "
            "and results, requests and chat, map). Shared pieces include routing (go_router), state management (Riverpod), "
            "theme and reusable widgets (glass-style cards, animations). Environment variables load API endpoints and keys "
            "without hard-coding secrets in source control.",
        ],
    )

    doc.add_heading("Data model and backend behaviour", level=2)
    add_paragraphs(
        doc,
        [
            "Supabase hosts PostgreSQL tables: profiles extend auth users; issues store diagnosis sessions; requests and "
            "request_responses model the help marketplace; chats and messages store conversation. Storage buckets hold "
            "profile images, issue media, voice clips, and chat images with policies that scope access by user or role.",
            "Database triggers create a profile row when a new auth user appears. Realtime publication on selected tables "
            "allows the client to subscribe to live inserts or updates for messages and request-related rows.",
        ],
    )

    doc.add_heading("External services", level=2)
    add_paragraphs(
        doc,
        [
            "The app may call Google’s generative AI API for natural-language explanations, and map/places APIs for "
            "location-aware features. These depend on valid API keys and acceptable use policies; they are optional "
            "failure points if keys are missing during development.",
        ],
    )

    doc.add_heading("Limitations and ethics", level=2)
    add_paragraphs(
        doc,
        [
            "AI output can be wrong or incomplete; the app should be presented as decision support, not a certified "
            "diagnosis. Road safety, privacy of location data, and fair dealing between users in help requests are "
            "important topics for any future production deployment beyond the graduation scope.",
        ],
    )

    doc.add_page_break()
    doc.add_heading("UML diagrams — how to read them", level=1)
    add_paragraphs(
        doc,
        [
            "The following pages show diagrams exported from PlantUML sources in docs/plantuml/. "
            "ERD = data structure and relationships. Use case = actors and main goals. Activity = ordered steps. "
            "Block = major system components. Sequence = message order between parts of the system over time.",
        ],
    )

    missing: list[str] = []

    for img_path, heading, paras in SECTIONS:
        doc.add_page_break()
        doc.add_heading(heading, level=1)
        for text in paras:
            doc.add_paragraph(text)

        if img_path.exists():
            doc.add_paragraph()
            pic_p = doc.add_paragraph()
            pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Fit page width ~6.2" for standard margins
            pic_p.add_run().add_picture(str(img_path), width=Inches(6.2))
        else:
            missing.append(str(img_path))
            doc.add_paragraph(f"[Image missing: {img_path.name} — run PlantUML on docs/plantuml/*.puml]")

    doc.add_page_break()
    doc.add_heading("Tools used (short list)", level=1)
    bullets = [
        "Flutter & Dart — mobile app",
        "Supabase — auth, PostgreSQL, storage, realtime",
        "Google Maps / Geolocator — location",
        "Google Generative AI (Gemini) — diagnosis",
        "Riverpod, go_router, Lottie, flutter_animate — UI and navigation",
        "PlantUML — diagrams (sources in docs/plantuml/)",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    if missing:
        doc.add_paragraph()
        doc.add_paragraph("Warning: some PNG files were missing when this file was built: " + "; ".join(missing))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(OUT))
        print(f"Wrote: {OUT}")
    except PermissionError:
        alt = OUT.with_name(f"{OUT.stem}_generated{OUT.suffix}")
        doc.save(str(alt))
        print(
            f"Could not overwrite {OUT.name} (close it in Word if open). "
            f"Wrote instead: {alt}"
        )


if __name__ == "__main__":
    main()
